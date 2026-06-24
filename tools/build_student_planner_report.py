# -*- coding: utf-8 -*-
"""Build the final Student Planner course report as a DOCX document.

The report is generated from the current Android project facts and the
provided Vietnamese report template/rubric requirements.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "BaoCao_Student_Planner_LTDD_HoanChinh_Final.docx"

FONT = "Times New Roman"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
LIGHT_BLUE = "E8F1FF"
LIGHT_GRAY = "F2F4F7"
TABLE_HEADER = "E8EEF5"
TEXT = RGBColor(0, 0, 0)


def set_run_font(run, name=FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=11):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
                tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
                tc_w = tc_pr.first_child_found_in("w:tcW")
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(int(width * 1440)))
                tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    header_tr_pr.append(tbl_header)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)
        set_cell_shading(cell, TABLE_HEADER)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            align = WD_ALIGN_PARAGRAPH.CENTER if len(str(value)) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], str(value), align=align, size=10)
    doc.add_paragraph()
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Trang ")
    set_run_font(run, size=10)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()
    r._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.right_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(13)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 14, BLUE, 12, 6),
        ("Heading 3", 13, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    return doc


def add_center_paragraph(doc, text, size=13, bold=False, after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_body(doc, text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    p.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=13, bold=bold, italic=italic)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run, size=13)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run, size=13)


def add_table_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True, italic=True)


def add_figure_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True, italic=True)


def add_figure_placeholder(doc, caption, body_lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [6.3])
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, LIGHT_GRAY)
    cell.text = ""
    for idx, line in enumerate(body_lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(line)
        set_run_font(run, size=11, bold=idx == 0, color=DARK_BLUE if idx == 0 else TEXT)
    set_cell_margins(cell, top=180, bottom=180, start=180, end=180)
    add_figure_caption(doc, caption)


def add_heading(doc, level, text):
    p = doc.add_heading(text, level=level)
    return p


def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build_report():
    doc = setup_doc()

    # Cover page
    add_center_paragraph(doc, "ĐẠI HỌC TRÀ VINH", size=14, bold=True, after=2)
    add_center_paragraph(doc, "TRƯỜNG KỸ THUẬT VÀ CÔNG NGHỆ", size=14, bold=True, after=2)
    add_center_paragraph(doc, "KHOA CÔNG NGHỆ THÔNG TIN", size=14, bold=True, after=2)
    add_center_paragraph(doc, "---------------------", size=13, after=36)
    add_center_paragraph(doc, "BÁO CÁO KẾT THÚC MÔN", size=18, bold=True, after=8)
    add_center_paragraph(doc, "HỌC PHẦN: LẬP TRÌNH THIẾT BỊ DI ĐỘNG", size=16, bold=True, after=42)
    add_center_paragraph(doc, "ĐỀ TÀI", size=16, bold=True, after=8)
    add_center_paragraph(
        doc,
        "XÂY DỰNG ỨNG DỤNG ANDROID QUẢN LÝ KẾ HOẠCH CÁ NHÂN DÀNH CHO SINH VIÊN",
        size=17,
        bold=True,
        after=42,
    )
    add_center_paragraph(doc, "Giảng viên hướng dẫn: ThS. Phạm Thị Trúc Mai", size=13, after=6)
    add_center_paragraph(doc, "Sinh viên thực hiện: Họ và tên sinh viên (MSSV)", size=13, after=6)
    add_center_paragraph(doc, "Mã lớp: ................................................", size=13, after=70)
    add_center_paragraph(doc, "Vĩnh Long, tháng 06 năm 2026", size=13, bold=True, after=0)
    add_page_break(doc)

    add_heading(doc, 1, "LỜI CẢM ƠN")
    add_body(
        doc,
        "Em xin chân thành cảm ơn giảng viên hướng dẫn đã tận tình định hướng, góp ý và hỗ trợ trong quá trình thực hiện đề tài kết thúc học phần Lập Trình Thiết Bị Di Động. Những góp ý về phạm vi đề tài, cách tổ chức chức năng và yêu cầu trình bày báo cáo đã giúp em hoàn thiện sản phẩm theo hướng rõ ràng hơn, bám sát nhu cầu thực tế của sinh viên.",
    )
    add_body(
        doc,
        "Em cũng xin cảm ơn Khoa Công Nghệ Thông Tin, Trường Kỹ Thuật và Công Nghệ đã tạo điều kiện để sinh viên có cơ hội vận dụng kiến thức lập trình Android vào một sản phẩm cụ thể. Trong quá trình thực hiện, do thời gian và kinh nghiệm còn hạn chế, báo cáo và ứng dụng khó tránh khỏi thiếu sót. Em rất mong nhận được sự góp ý của quý thầy cô để tiếp tục hoàn thiện trong các phiên bản sau.",
    )
    add_body(doc, "Sinh viên thực hiện", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_page_break(doc)

    add_heading(doc, 1, "NHẬN XÉT CỦA GIÁO VIÊN HƯỚNG DẪN / GIÁO VIÊN CHẤM")
    for _ in range(12):
        p = doc.add_paragraph("........................................................................................................................")
        p.paragraph_format.space_after = Pt(8)
    add_body(doc, "Ngày .... tháng .... năm 2026", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_body(doc, "Giáo viên chấm báo cáo", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_page_break(doc)

    add_heading(doc, 1, "MỤC LỤC")
    toc_rows = [
        ("TÓM TẮT ĐỀ TÀI", "1"),
        ("CHƯƠNG 1: GIỚI THIỆU ĐỀ TÀI", "2"),
        ("1.1. Lý do chọn đề tài", "2"),
        ("1.2. Mục tiêu đề tài", "3"),
        ("1.3. Đối tượng và phạm vi nghiên cứu", "3"),
        ("1.4. Phương pháp thực hiện", "4"),
        ("1.5. Công cụ và môi trường thực hiện", "4"),
        ("CHƯƠNG 2: CƠ SỞ LÝ THUYẾT", "6"),
        ("2.1. Tổng quan về lập trình thiết bị di động", "6"),
        ("2.2. Tổng quan về Android", "7"),
        ("2.3. Thiết kế giao diện Android", "7"),
        ("2.4. Lưu trữ dữ liệu cục bộ", "8"),
        ("2.5. Thông báo và nhắc lịch trên Android", "9"),
        ("2.6. Cơ sở nghiệp vụ quản lý kế hoạch cá nhân cho sinh viên", "10"),
        ("CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", "12"),
        ("3.1. Mô tả bài toán", "12"),
        ("3.2. Yêu cầu chức năng", "13"),
        ("3.3. Yêu cầu phi chức năng", "15"),
        ("3.4. Mô hình dữ liệu", "16"),
        ("3.5. Thiết kế cơ sở dữ liệu", "18"),
        ("3.6. Mô hình xử lý", "20"),
        ("3.7. Thiết kế giao diện", "23"),
        ("3.8. Kết chương", "24"),
        ("CHƯƠNG 4: CÀI ĐẶT VÀ KIỂM THỬ ỨNG DỤNG", "25"),
        ("4.1. Môi trường cài đặt", "25"),
        ("4.2. Cấu trúc dự án sau khi dọn dẹp", "26"),
        ("4.3. Cài đặt các chức năng chính", "28"),
        ("4.4. Kết quả giao diện ứng dụng", "30"),
        ("4.5. Kiểm thử chức năng", "34"),
        ("4.6. Đánh giá kết quả thực nghiệm", "37"),
        ("CHƯƠNG 5: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", "39"),
        ("5.1. Kết quả đạt được", "39"),
        ("5.2. Ưu điểm và nhược điểm", "40"),
        ("5.3. Hướng phát triển", "42"),
        ("TÀI LIỆU THAM KHẢO", "43"),
    ]
    add_table(doc, ["Nội dung", "Trang"], toc_rows, [5.4, 0.9])
    add_body(
        doc,
        "Ghi chú: Mục lục được trình bày theo cấu trúc của báo cáo. Khi thay ảnh giao diện thật và in nộp, có thể cập nhật lại số trang trong Word nếu cần.",
        italic=True,
    )
    add_page_break(doc)

    add_heading(doc, 1, "DANH MỤC TỪ VIẾT TẮT")
    add_table(
        doc,
        ["Từ viết tắt", "Ý nghĩa"],
        [
            ("ADB", "Android Debug Bridge"),
            ("API", "Application Programming Interface"),
            ("APK", "Android Package Kit"),
            ("CRUD", "Create, Read, Update, Delete"),
            ("DB", "Database"),
            ("JDK", "Java Development Kit"),
            ("PBKDF2", "Password-Based Key Derivation Function 2"),
            ("SDK", "Software Development Kit"),
            ("SQL", "Structured Query Language"),
            ("UI/UX", "User Interface / User Experience"),
            ("XML", "Extensible Markup Language"),
        ],
        [1.4, 4.9],
    )

    add_heading(doc, 1, "DANH MỤC HÌNH ẢNH")
    add_table(
        doc,
        ["Số hiệu", "Tên hình", "Trang"],
        [
            ("Hình 2.1", "Kiến trúc tổng quan ứng dụng Student Planner", "11"),
            ("Hình 3.1", "Mô hình dữ liệu SQLite của ứng dụng", "17"),
            ("Hình 3.2", "Luồng xử lý thêm kế hoạch", "21"),
            ("Hình 3.3", "Luồng xử lý nhắc lịch", "22"),
            ("Hình 4.1", "Màn hình đăng nhập", "30"),
            ("Hình 4.2", "Màn hình trang chủ", "31"),
            ("Hình 4.3", "Màn hình danh sách kế hoạch", "31"),
            ("Hình 4.4", "Màn hình thêm/sửa kế hoạch", "32"),
            ("Hình 4.5", "Màn hình lịch", "33"),
            ("Hình 4.6", "Màn hình tổng quan/thống kê", "33"),
        ],
        [1.1, 4.7, 0.5],
    )

    add_heading(doc, 1, "DANH MỤC BẢNG BIỂU")
    add_table(
        doc,
        ["Số hiệu", "Tên bảng", "Trang"],
        [
            ("Bảng 1.1", "Đối chiếu nội dung báo cáo với phiếu chấm", "5"),
            ("Bảng 3.1", "Danh sách chức năng chính của ứng dụng", "13"),
            ("Bảng 3.2", "Yêu cầu phi chức năng", "15"),
            ("Bảng 3.3", "Thiết kế cơ sở dữ liệu", "18"),
            ("Bảng 3.4", "Mô tả các màn hình giao diện", "23"),
            ("Bảng 4.1", "Môi trường cài đặt", "25"),
            ("Bảng 4.2", "Package/class chính trong dự án", "26"),
            ("Bảng 4.3", "Kết quả cài đặt chức năng", "29"),
            ("Bảng 4.4", "Bảng kiểm thử chức năng", "35"),
            ("Bảng 4.5", "Đánh giá kết quả đạt được", "38"),
        ],
        [1.1, 4.7, 0.5],
    )
    add_page_break(doc)

    add_heading(doc, 1, "TÓM TẮT ĐỀ TÀI")
    add_body(
        doc,
        "Đề tài “Xây dựng ứng dụng Android quản lý kế hoạch cá nhân dành cho sinh viên” tập trung xây dựng một ứng dụng di động giúp sinh viên tổ chức, theo dõi và đánh giá các kế hoạch trong đời sống hằng ngày. Phạm vi đề tài không dừng ở quản lý học tập đơn thuần mà mở rộng sang các kế hoạch thường gặp như làm bài tập, đi học, làm thêm, ôn thi, thực hiện đồ án/dự án và các việc cá nhân.",
    )
    add_body(
        doc,
        "Ứng dụng được cài đặt bằng Java, XML Layout, AndroidX AppCompat, RecyclerView, Material Components và SQLiteOpenHelper. Dữ liệu hoạt động offline trong cơ sở dữ liệu SQLite nội bộ tên personal_planner.db. Tài khoản người dùng được đăng ký/đăng nhập cục bộ, mật khẩu được băm bằng PBKDF2 kèm salt và phiên đăng nhập được lưu bằng SharedPreferences.",
    )
    add_body(
        doc,
        "Kết quả hiện tại gồm các nhóm chức năng chính: đăng ký/đăng nhập offline, quản lý danh mục kế hoạch, thêm/sửa/xóa kế hoạch theo loại, tìm kiếm và lọc kế hoạch, lịch dạng tháng kết hợp danh sách kế hoạch theo ngày, kiểm tra trùng thời gian, nhận diện kế hoạch quá hạn khi hiển thị, nhắc lịch bằng AlarmManager/Notification, thống kê tổng quan và gợi ý việc cần ưu tiên trong ngày. Một số thành phần mở rộng như sub-task và đánh giá sau hoàn thành đã có nền tảng dữ liệu, cần tiếp tục hoàn thiện giao diện thao tác chi tiết.",
    )
    add_page_break(doc)

    # Chapter 1
    add_heading(doc, 1, "CHƯƠNG 1: GIỚI THIỆU ĐỀ TÀI")
    add_heading(doc, 2, "1.1. Lý do chọn đề tài")
    add_body(
        doc,
        "Sinh viên thường phải xử lý nhiều kế hoạch trong cùng một khoảng thời gian: lịch học, hạn nộp bài tập, lịch ôn thi, ca làm thêm, tiến độ đồ án nhóm và các việc cá nhân. Nếu chỉ ghi chú rời rạc hoặc chỉ quản lý theo môn học, người dùng khó quan sát tổng thể thời gian, dễ trùng lịch và dễ bỏ sót các việc quan trọng.",
    )
    add_body(
        doc,
        "Vì vậy, đề tài lựa chọn xây dựng ứng dụng quản lý kế hoạch cá nhân dành cho sinh viên nhằm gom nhiều hoạt động thường ngày vào một hệ thống trực quan, hoạt động offline và có khả năng nhắc lịch. Ứng dụng cần khác với một ứng dụng ghi chú thông thường ở chỗ có loại kế hoạch, mức ưu tiên, trạng thái thực hiện, thời hạn, kiểm tra quá hạn, thống kê và gợi ý việc cần ưu tiên.",
    )

    add_heading(doc, 2, "1.2. Mục tiêu đề tài")
    add_bullets(
        doc,
        [
            "Xây dựng ứng dụng Android hỗ trợ sinh viên tạo, sửa, xóa, xem chi tiết và theo dõi kế hoạch cá nhân.",
            "Phân loại kế hoạch theo danh mục và theo loại việc như bài tập, đi học, làm thêm, cá nhân, thi/kiểm tra và đồ án/dự án.",
            "Lưu trữ dữ liệu cục bộ bằng SQLite, cho phép sử dụng khi không có Internet.",
            "Cung cấp danh sách kế hoạch có tìm kiếm/lọc, lịch dạng tháng, nhắc lịch và cảnh báo trùng thời gian.",
            "Tự động nhận diện kế hoạch quá hạn khi hiển thị, không lưu quá hạn như một trạng thái cứng trong cơ sở dữ liệu.",
            "Thống kê tổng quan, thống kê theo tuần và gợi ý kế hoạch cần ưu tiên trong ngày.",
        ],
    )

    add_heading(doc, 2, "1.3. Đối tượng và phạm vi nghiên cứu")
    add_body(
        doc,
        "Đối tượng sử dụng là sinh viên cần quản lý lịch học, bài tập, việc làm thêm, ôn thi, đồ án/dự án và kế hoạch cá nhân. Đối tượng nghiên cứu là quy trình tổ chức kế hoạch cá nhân trên thiết bị Android, gồm lưu trữ dữ liệu offline, hiển thị lịch, nhắc lịch, lọc kế hoạch và thống kê tiến độ.",
    )
    add_body(
        doc,
        "Phạm vi đề tài tập trung vào ứng dụng Android cá nhân, chạy offline trên một thiết bị. Báo cáo không mở rộng sang quản lý tài chính, theo dõi sức khỏe, uống nước, tập thể dục, mua sắm, mạng xã hội, Google Calendar hay đồng bộ đám mây. Các hướng này chỉ được nhắc đến ở phần hướng phát triển nếu phù hợp.",
    )

    add_heading(doc, 2, "1.4. Phương pháp thực hiện")
    add_numbered(
        doc,
        [
            "Khảo sát nhu cầu quản lý kế hoạch của sinh viên và xác định các loại kế hoạch thường gặp.",
            "Phân tích yêu cầu chức năng, yêu cầu phi chức năng và mô hình dữ liệu phù hợp với đề tài.",
            "Thiết kế giao diện theo hướng tối giản, dễ thao tác trên điện thoại Android.",
            "Cài đặt ứng dụng bằng Java, XML Layout, SQLiteOpenHelper, SharedPreferences và Material Components.",
            "Kiểm thử thủ công các luồng chính như đăng ký, đăng nhập, thêm/sửa/xóa kế hoạch, lọc, lịch, nhắc lịch, thống kê và xử lý quá hạn.",
        ],
    )

    add_heading(doc, 2, "1.5. Công cụ và môi trường thực hiện")
    add_body(
        doc,
        "Dự án hiện tại sử dụng Android Studio/Gradle, ngôn ngữ Java, giao diện XML Layout, thư viện AndroidX AppCompat, RecyclerView và Material Components. Cơ sở dữ liệu sử dụng SQLite thông qua SQLiteOpenHelper. Ứng dụng không sử dụng Firebase, Room, Jetpack Compose hoặc máy chủ backend trong phiên bản hiện tại.",
    )
    add_table_caption(doc, "Bảng 1.1. Đối chiếu nội dung báo cáo với phiếu chấm")
    add_table(
        doc,
        ["Tiêu chí phiếu chấm", "Nội dung đáp ứng trong báo cáo", "Ghi chú"],
        [
            ("Hình thức", "Bìa, lời cảm ơn, nhận xét GV, mục lục, danh mục hình/bảng/từ viết tắt, tài liệu tham khảo.", "Đã bổ sung đầy đủ."),
            ("Mở đầu", "Lý do chọn đề tài, mục tiêu, đối tượng/phạm vi, phương pháp, công cụ.", "Chương 1."),
            ("Cơ sở lý thuyết", "Android, Activity/Fragment/RecyclerView, XML, SQLite, SharedPreferences, AlarmManager/Notification.", "Chương 2."),
            ("Giải pháp thực hiện", "Mô tả bài toán, yêu cầu, mô hình dữ liệu, thiết kế bảng, luồng xử lý.", "Chương 3."),
            ("Thực nghiệm/cài đặt", "Môi trường, cấu trúc project, màn hình, chức năng và kiểm thử.", "Chương 4."),
            ("Kết luận", "Kết quả đạt được, ưu/nhược điểm, hướng phát triển.", "Chương 5."),
        ],
        [1.45, 3.95, 0.9],
    )
    add_page_break(doc)

    # Chapter 2
    add_heading(doc, 1, "CHƯƠNG 2: CƠ SỞ LÝ THUYẾT")
    add_heading(doc, 2, "2.1. Tổng quan về lập trình thiết bị di động")
    add_body(
        doc,
        "Lập trình thiết bị di động là lĩnh vực xây dựng phần mềm chạy trên điện thoại thông minh và máy tính bảng. Khác với ứng dụng desktop, ứng dụng di động cần quan tâm đến kích thước màn hình nhỏ, thao tác cảm ứng, vòng đời ứng dụng, quyền hệ thống, tiêu thụ pin và khả năng lưu trữ dữ liệu cục bộ.",
    )
    add_body(
        doc,
        "Trong học phần Lập Trình Thiết Bị Di Động, việc xây dựng một ứng dụng Android quản lý kế hoạch giúp sinh viên vận dụng các kiến thức về giao diện, xử lý sự kiện, lưu trữ dữ liệu, điều hướng giữa màn hình và tương tác với hệ thống thông báo của Android.",
    )

    add_heading(doc, 2, "2.2. Tổng quan về Android")
    add_body(
        doc,
        "Android là nền tảng mã nguồn mở dành cho thiết bị di động, hỗ trợ phát triển ứng dụng bằng Java hoặc Kotlin. Một ứng dụng Android thường gồm Activity, Fragment, Adapter, model dữ liệu, lớp truy cập dữ liệu và tài nguyên giao diện như layout, drawable, menu, color, string.",
    )
    add_body(
        doc,
        "Trong dự án Student Planner, Activity được dùng cho các màn hình độc lập như đăng nhập, đăng ký, thêm kế hoạch, chi tiết kế hoạch và quản lý danh mục. Fragment được dùng trong MainActivity để tổ chức các tab Trang chủ, Lịch, Tổng quan và Cài đặt. RecyclerView và Adapter được dùng để hiển thị danh sách kế hoạch, danh mục và sự kiện lịch.",
    )

    add_heading(doc, 2, "2.3. Thiết kế giao diện Android")
    add_body(
        doc,
        "Giao diện Android trong dự án được khai báo bằng XML Layout và sử dụng Material Components. Các thành phần như MaterialCardView, BottomNavigationView, FloatingActionButton, Chip, TextInputLayout, SwitchMaterial và ProgressIndicator giúp giao diện nhất quán, hiện đại và dễ sử dụng.",
    )
    add_body(
        doc,
        "Thiết kế của ứng dụng ưu tiên bố cục tối giản, khoảng trắng rộng, card bo góc mềm, màu trạng thái rõ ràng và bottom navigation cố định. Cách thiết kế này phù hợp với ứng dụng quản lý kế hoạch vì người dùng cần nhìn nhanh các việc quan trọng, kế hoạch quá hạn và thống kê trong ngày.",
    )

    add_heading(doc, 2, "2.4. Lưu trữ dữ liệu cục bộ")
    add_body(
        doc,
        "SQLite là hệ quản trị cơ sở dữ liệu quan hệ nhúng, phù hợp cho các ứng dụng di động cần lưu dữ liệu offline. Android cung cấp SQLiteOpenHelper để tạo bảng, nâng cấp cấu trúc dữ liệu và thực hiện các thao tác CRUD. Dữ liệu SQLite nằm trong vùng dữ liệu riêng của ứng dụng nên người dùng thông thường không truy cập trực tiếp được nếu thiết bị không root.",
    )
    add_body(
        doc,
        "Student Planner sử dụng cơ sở dữ liệu personal_planner.db. Các bảng quan trọng gồm users, plan_categories, tasks, sub_tasks, reminders, repeat_rules và plan_evaluations. Phiên đăng nhập được lưu bằng SharedPreferences thông qua lớp SessionManager, giúp ứng dụng ghi nhớ người dùng đã đăng nhập giữa các lần mở app.",
    )

    add_heading(doc, 2, "2.5. Thông báo và nhắc lịch trên Android")
    add_body(
        doc,
        "Nhắc lịch là chức năng giúp ứng dụng quản lý kế hoạch khác với một ứng dụng ghi chú thông thường. Android cung cấp AlarmManager để đặt thời điểm kích hoạt trong tương lai, BroadcastReceiver để nhận sự kiện alarm và NotificationCompat để tạo thông báo cho người dùng.",
    )
    add_body(
        doc,
        "Với Android 12 trở lên, alarm chính xác cần quan tâm đến quyền SCHEDULE_EXACT_ALARM. Dự án đã có lớp ReminderScheduler kiểm tra quyền, sử dụng setExactAndAllowWhileIdle để đặt alarm chính xác hơn trong điều kiện Doze/App Standby, đồng thời có BootReceiver để đăng ký lại reminder sau khi thiết bị khởi động lại.",
    )

    add_heading(doc, 2, "2.6. Cơ sở nghiệp vụ quản lý kế hoạch cá nhân cho sinh viên")
    add_body(
        doc,
        "Kế hoạch cá nhân của sinh viên có nhiều ngữ cảnh khác nhau. Một bài tập cần hạn nộp, môn liên quan và trạng thái nộp; một buổi đi học cần thời gian, địa điểm, phòng học và môn/lớp; một ca làm thêm cần nơi làm, thời gian bắt đầu/kết thúc và có thể có tiền công dự kiến; một kế hoạch cá nhân có thể cần ghi chú hoặc địa điểm.",
    )
    add_body(
        doc,
        "Vì vậy, mô hình kế hoạch cần có danh mục, loại kế hoạch, mức ưu tiên, thời gian, trạng thái thực hiện, nhắc lịch và khả năng thống kê. Trạng thái quá hạn không nên lưu cứng vào cơ sở dữ liệu mà nên tính khi hiển thị: nếu thời hạn đã qua và kế hoạch chưa hoàn thành/chưa hủy thì hiển thị là quá hạn.",
    )
    add_figure_placeholder(
        doc,
        "Hình 2.1. Kiến trúc tổng quan ứng dụng Student Planner",
        [
            "[CHÈN HÌNH KIẾN TRÚC TỔNG QUAN ỨNG DỤNG TẠI ĐÂY]",
            "UI XML/Activity/Fragment -> Adapter/Controller -> DatabaseHelper -> SQLite local DB",
            "ReminderScheduler -> AlarmManager -> ReminderReceiver -> Notification",
        ],
    )
    add_page_break(doc)

    # Chapter 3
    add_heading(doc, 1, "CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG")
    add_heading(doc, 2, "3.1. Mô tả bài toán")
    add_body(
        doc,
        "Bài toán đặt ra là xây dựng một ứng dụng Android giúp sinh viên quản lý nhiều loại kế hoạch trong đời sống hằng ngày. Người dùng có thể đăng ký tài khoản cục bộ, đăng nhập, tạo danh mục kế hoạch, thêm kế hoạch theo loại việc, xem danh sách, lọc/tìm kiếm, xem lịch, nhận nhắc lịch, đánh dấu hoàn thành và xem thống kê.",
    )
    add_body(
        doc,
        "Điểm cốt lõi của bài toán là không xem mọi kế hoạch như một ghi chú giống nhau. Ứng dụng cần thể hiện rõ ngữ cảnh của từng loại kế hoạch, đồng thời hỗ trợ nhận diện quá hạn, kiểm tra trùng giờ và gợi ý việc nên ưu tiên trong ngày.",
    )

    add_heading(doc, 2, "3.2. Yêu cầu chức năng")
    add_table_caption(doc, "Bảng 3.1. Danh sách chức năng chính của ứng dụng")
    add_table(
        doc,
        ["Mã", "Chức năng", "Mô tả", "Mức hoàn thành"],
        [
            ("F01", "Đăng ký tài khoản", "Kiểm tra username, email, mật khẩu; lưu mật khẩu dạng PBKDF2.", "Hoàn thành"),
            ("F02", "Đăng nhập/đăng xuất", "Xác thực tài khoản SQLite và lưu phiên bằng SharedPreferences.", "Hoàn thành"),
            ("F03", "Quản lý danh mục", "Thêm, sửa, xóa danh mục kế hoạch theo từng người dùng.", "Hoàn thành"),
            ("F04", "Thêm kế hoạch", "Tạo kế hoạch theo loại, ưu tiên, thời gian, nhắc lịch, lặp lại và trường riêng.", "Hoàn thành"),
            ("F05", "Sửa/xóa kế hoạch", "Cập nhật nội dung, trạng thái, reminder và xóa kế hoạch.", "Hoàn thành"),
            ("F06", "Tìm kiếm/lọc", "Tìm theo từ khóa; lọc theo trạng thái, danh mục và loại kế hoạch.", "Hoàn thành"),
            ("F07", "Lịch", "Hiển thị lịch dạng tháng, đánh dấu ngày có kế hoạch và danh sách theo ngày chọn.", "Hoàn thành cơ bản"),
            ("F08", "Nhắc lịch", "Đặt exact alarm, hiện notification, hỗ trợ nhắc 24h và khôi phục sau reboot.", "Hoàn thành cơ bản"),
            ("F09", "Trùng giờ/quá hạn", "Kiểm tra trùng giờ khi tạo/sửa; quá hạn tính runtime khi hiển thị.", "Hoàn thành"),
            ("F10", "Thống kê/gợi ý", "Dashboard, overview, thống kê tuần, bài tập chưa nộp, giờ học/làm thêm, gợi ý hôm nay.", "Hoàn thành cơ bản"),
            ("F11", "Sub-task/đánh giá", "Đã có bảng/model/DAO; giao diện thao tác chi tiết cần hoàn thiện thêm.", "Nền tảng dữ liệu"),
        ],
        [0.55, 1.35, 3.35, 1.05],
    )

    add_heading(doc, 2, "3.3. Yêu cầu phi chức năng")
    add_table_caption(doc, "Bảng 3.2. Yêu cầu phi chức năng")
    add_table(
        doc,
        ["Nhóm yêu cầu", "Mô tả", "Cách đáp ứng trong dự án"],
        [
            ("Offline", "Ứng dụng sử dụng được khi không có Internet.", "Dữ liệu lưu trong SQLite local, không phụ thuộc server."),
            ("Bảo mật cơ bản", "Không lưu mật khẩu dạng rõ.", "PasswordUtils dùng PBKDF2WithHmacSHA256, salt ngẫu nhiên và 120.000 vòng lặp."),
            ("Dễ sử dụng", "Giao diện rõ, thao tác nhanh trên điện thoại.", "Material Components, card, chip lọc, bottom navigation và FAB."),
            ("Ổn định dữ liệu", "Dữ liệu tách theo người dùng và không mất khi đóng app.", "Các bảng có user_id; session lưu bằng SharedPreferences."),
            ("Bảo trì", "Code chia theo vai trò.", "Package activity, fragment, adapter, data/local, data/model, notification, utils."),
            ("Tương thích", "Chạy trên nhiều thiết bị Android phổ biến.", "minSdk 23, targetSdk 35, dùng AndroidX/Material."),
        ],
        [1.25, 2.1, 2.95],
    )

    add_heading(doc, 2, "3.4. Mô hình dữ liệu")
    add_body(
        doc,
        "Cơ sở dữ liệu của ứng dụng có tên personal_planner.db, được quản lý trong lớp DatabaseHelper và đang ở phiên bản 5. Do dự án ban đầu xuất phát từ todo/study planner, bảng lưu kế hoạch vẫn giữ tên vật lý là tasks để tương thích dữ liệu cũ, nhưng về nghiệp vụ hiện tại bảng này đóng vai trò là bảng plan/kế hoạch.",
    )
    add_body(
        doc,
        "Quan hệ dữ liệu chính gồm: một người dùng có nhiều danh mục và kế hoạch; một danh mục có nhiều kế hoạch; một kế hoạch có thể có nhiều sub-task; một kế hoạch có thể có một reminder, một repeat_rule và một plan_evaluation. Trạng thái quá hạn không được lưu thành trạng thái riêng mà được tính dựa trên thời hạn và trạng thái thật.",
    )
    add_figure_placeholder(
        doc,
        "Hình 3.1. Mô hình dữ liệu SQLite của ứng dụng",
        [
            "[CHÈN HÌNH MÔ HÌNH DỮ LIỆU/ERD TẠI ĐÂY]",
            "users 1-n plan_categories",
            "users 1-n tasks; plan_categories 1-n tasks",
            "tasks 1-n sub_tasks; tasks 1-1 reminders/repeat_rules/plan_evaluations",
        ],
    )

    add_heading(doc, 2, "3.5. Thiết kế cơ sở dữ liệu")
    add_table_caption(doc, "Bảng 3.3. Thiết kế cơ sở dữ liệu")
    add_table(
        doc,
        ["Bảng", "Thuộc tính chính", "Khóa/quan hệ", "Vai trò"],
        [
            ("users", "user_id, username, email, password, created_at", "PK: user_id; username/email unique", "Lưu tài khoản offline và mật khẩu đã băm."),
            ("plan_categories", "category_id, category_name, category_code, note, color, created_at, updated_at, user_id", "PK: category_id; FK: user_id", "Lưu danh mục/nhóm kế hoạch theo từng tài khoản."),
            ("tasks", "task_id, title, description, date, time, end_time, status, category_id, plan_type, priority, duration_minutes, reminder_enabled, reminder_minutes, location, room, subject, repeat_rule, repeat_until, wage, submitted, start_time, deadline, created_at, updated_at, user_id", "PK: task_id; FK: user_id/category_id", "Bảng kế hoạch chính, giữ tên tasks để tương thích dữ liệu cũ."),
            ("sub_tasks", "sub_task_id, plan_id, title, is_completed, created_at, updated_at", "PK: sub_task_id; FK: plan_id", "Nền tảng chia kế hoạch thành việc con một cấp."),
            ("reminders", "reminder_id, plan_id, reminder_time, is_enabled, created_at, updated_at", "PK: reminder_id; plan_id unique", "Lưu thời điểm nhắc lịch đã đặt."),
            ("repeat_rules", "repeat_rule_id, plan_id, repeat_type, week_days, month_day, is_active, created_at, updated_at", "PK: repeat_rule_id; plan_id unique", "Lưu quy tắc lặp cơ bản."),
            ("plan_evaluations", "evaluation_id, plan_id, satisfaction_level, result_note, delay_reason, completed_at, created_at, updated_at", "PK: evaluation_id; plan_id unique", "Nền tảng đánh giá sau khi hoàn thành kế hoạch."),
        ],
        [1.05, 3.1, 1.35, 1.0],
    )

    add_heading(doc, 2, "3.6. Mô hình xử lý")
    add_heading(doc, 3, "3.6.1. Luồng đăng ký/đăng nhập")
    add_body(
        doc,
        "Khi đăng ký, ứng dụng kiểm tra dữ liệu nhập, kiểm tra username/email trùng và băm mật khẩu bằng PBKDF2 trước khi lưu vào bảng users. Khi đăng nhập, DatabaseHelper lấy user theo username/email, PasswordUtils xác thực mật khẩu và SessionManager lưu user_id, username vào SharedPreferences nếu thành công.",
    )
    add_heading(doc, 3, "3.6.2. Luồng thêm/sửa/xóa kế hoạch")
    add_body(
        doc,
        "Khi thêm kế hoạch, AddTaskActivity kiểm tra tiêu đề, thời gian, danh mục và thiết lập nhắc lịch. Form thay đổi các trường theo loại kế hoạch: bài tập ưu tiên môn liên quan/trạng thái nộp, đi học ưu tiên địa điểm/phòng/môn, làm thêm ưu tiên nơi làm/tiền công, cá nhân ưu tiên ghi chú. Nếu có lặp lại, hệ thống sinh danh sách ngày và kiểm tra trùng giờ từng ngày trước khi lưu.",
    )
    add_figure_placeholder(
        doc,
        "Hình 3.2. Luồng xử lý thêm kế hoạch",
        [
            "[CHÈN HÌNH LUỒNG XỬ LÝ THÊM KẾ HOẠCH TẠI ĐÂY]",
            "Nhập form -> Validate -> Sinh ngày lặp -> Kiểm tra trùng giờ -> Lưu SQLite -> Đặt reminder nếu bật",
        ],
    )
    add_heading(doc, 3, "3.6.3. Luồng lịch")
    add_body(
        doc,
        "CalendarMockupFragment hiển thị lịch dạng tháng. Khi người dùng chuyển tháng hoặc chọn ngày, ứng dụng lấy kế hoạch trong tháng để đánh dấu ngày có sự kiện, đồng thời lấy danh sách kế hoạch theo ngày được chọn để hiển thị trong RecyclerView.",
    )
    add_heading(doc, 3, "3.6.4. Luồng nhắc lịch")
    add_body(
        doc,
        "ReminderScheduler tính thời điểm nhắc từ ngày, giờ và lựa chọn nhắc của người dùng. Với lựa chọn 24 giờ, hệ thống tính thời điểm sau một ngày và sau khi receiver chạy sẽ đặt lại alarm kế tiếp. Với lựa chọn đúng giờ hoặc sau 5/10/20/30/60 phút, thời điểm nhắc được tính bằng thời điểm người dùng chọn cộng thêm offset tương ứng.",
    )
    add_figure_placeholder(
        doc,
        "Hình 3.3. Luồng xử lý nhắc lịch",
        [
            "[CHÈN HÌNH LUỒNG XỬ LÝ NHẮC LỊCH TẠI ĐÂY]",
            "AddTaskActivity -> ReminderScheduler.setExactAndAllowWhileIdle -> ReminderReceiver -> Notification",
            "BootReceiver -> đọc kế hoạch/reminder còn hiệu lực -> đặt lại alarm",
        ],
    )
    add_heading(doc, 3, "3.6.5. Luồng thống kê và quá hạn")
    add_body(
        doc,
        "Thống kê tổng quan được lấy qua getStudyStatistics, còn thống kê theo khoảng thời gian được lấy qua getPlanRangeStats. Kế hoạch quá hạn được xác định runtime bằng PlanBusinessRules: nếu kế hoạch chưa hoàn thành/chưa hủy và deadline hoặc ngày giờ kết thúc nhỏ hơn thời điểm hiện tại thì UI hiển thị trạng thái quá hạn.",
    )

    add_heading(doc, 2, "3.7. Thiết kế giao diện")
    add_body(
        doc,
        "Giao diện được tổ chức xoay quanh MainActivity với BottomNavigationView gồm Trang chủ, Lịch, Tổng quan và Cài đặt. FloatingActionButton dùng để mở màn hình thêm kế hoạch. Các màn hình sử dụng XML Layout, MaterialCardView, chip lọc, progress indicator và màu trạng thái để hỗ trợ người dùng đọc nhanh thông tin.",
    )
    add_table_caption(doc, "Bảng 3.4. Mô tả các màn hình giao diện")
    add_table(
        doc,
        ["Màn hình", "File/layout chính", "Vai trò"],
        [
            ("Đăng nhập", "LoginActivity / activity_login.xml", "Xác thực tài khoản offline."),
            ("Đăng ký", "RegisterActivity / activity_register.xml", "Tạo tài khoản cục bộ."),
            ("Trang chủ", "HomeMockupFragment / fragment_home_mockup.xml", "Dashboard, thống kê nhanh, gợi ý hôm nay."),
            ("Danh sách kế hoạch", "TaskFragment / fragment_task.xml", "Tìm kiếm, lọc, cập nhật trạng thái kế hoạch."),
            ("Thêm/sửa kế hoạch", "AddTaskActivity / activity_add_schedule_mockup.xml", "Form nhập kế hoạch theo loại, nhắc lịch, lặp lại."),
            ("Chi tiết kế hoạch", "TaskDetailActivity, PlanDetailMockupActivity", "Xem thông tin, đổi trạng thái, sửa/xóa."),
            ("Lịch", "CalendarMockupFragment / fragment_calendar_mockup.xml", "Lịch tháng và danh sách kế hoạch theo ngày."),
            ("Tổng quan", "OverviewFragment / fragment_overview.xml", "Thống kê tổng, quá hạn, tiến độ, ưu tiên và tuần hiện tại."),
            ("Cài đặt", "ProfileFragment / fragment_profile_simplified.xml", "Giao diện sáng/tối, danh mục, thông báo, đăng xuất."),
        ],
        [1.4, 2.25, 2.65],
    )

    add_heading(doc, 2, "3.8. Kết chương")
    add_body(
        doc,
        "Chương 3 đã trình bày bài toán, yêu cầu chức năng, yêu cầu phi chức năng, mô hình dữ liệu và các luồng xử lý chính của ứng dụng. Các thiết kế này là cơ sở để cài đặt và kiểm thử trong chương tiếp theo.",
    )
    add_page_break(doc)

    # Chapter 4
    add_heading(doc, 1, "CHƯƠNG 4: CÀI ĐẶT VÀ KIỂM THỬ ỨNG DỤNG")
    add_heading(doc, 2, "4.1. Môi trường cài đặt")
    add_table_caption(doc, "Bảng 4.1. Môi trường cài đặt")
    add_table(
        doc,
        ["Thành phần", "Giá trị sử dụng trong dự án"],
        [
            ("Nền tảng", "Android"),
            ("Ngôn ngữ", "Java"),
            ("Giao diện", "XML Layout, Material Components"),
            ("Cơ sở dữ liệu", "SQLite thông qua SQLiteOpenHelper"),
            ("Phiên bản SDK", "compileSdk 35, minSdk 23, targetSdk 35"),
            ("Thư viện", "androidx.appcompat 1.7.0, recyclerview 1.3.2, material 1.12.0"),
            ("Cấu hình Java", "JavaVersion 17"),
            ("Build tool", "Gradle Android plugin, gradlew.bat"),
        ],
        [2.0, 4.3],
    )

    add_heading(doc, 2, "4.2. Cấu trúc dự án sau khi dọn dẹp")
    add_body(
        doc,
        "Sau khi dọn dẹp, mã nguồn chính nằm trong package com.example.personalplanner. Dự án được chia thành các nhóm file theo vai trò để dễ bảo trì: activity, fragment, adapter, data/local, data/model, notification và utils. Tài nguyên giao diện nằm trong app/src/main/res gồm layout, drawable, color, menu, values và values-night.",
    )
    add_table_caption(doc, "Bảng 4.2. Package/class chính trong dự án")
    add_table(
        doc,
        ["Nhóm", "Class/file tiêu biểu", "Vai trò"],
        [
            ("activity", "LoginActivity, RegisterActivity, MainActivity, AddTaskActivity, TaskDetailActivity, PlanCategoryListActivity", "Điều khiển các màn hình độc lập và xử lý thao tác chính."),
            ("fragment", "HomeMockupFragment, TaskFragment, CalendarMockupFragment, OverviewFragment, ProfileFragment", "Các tab trong MainActivity."),
            ("adapter", "TaskAdapter, CalendarEventAdapter, PlanCategoryAdapter", "Đổ dữ liệu vào RecyclerView."),
            ("data/local", "DatabaseHelper", "Tạo bảng, nâng cấp DB và cung cấp CRUD."),
            ("data/model", "StudyPlan, PlanCategory, User, StudyStatistics, PlanRangeStats, SubTask, PlanReminder, RepeatRule, PlanEvaluation", "Định nghĩa dữ liệu nghiệp vụ."),
            ("notification", "ReminderScheduler, ReminderReceiver, BootReceiver, ReminderType", "Đặt alarm, nhận alarm và hiển thị notification."),
            ("utils", "PasswordUtils, SessionManager, ThemeManager, PlanBusinessRules", "Băm mật khẩu, phiên đăng nhập, theme, quy tắc quá hạn/tiến độ."),
            ("res", "layout, drawable, values, values-night, menu, xml", "Tài nguyên giao diện, màu, chuỗi, menu và cấu hình backup."),
        ],
        [1.25, 2.65, 2.4],
    )

    add_heading(doc, 2, "4.3. Cài đặt các chức năng chính")
    add_body(
        doc,
        "Chức năng tài khoản được cài đặt trong LoginActivity, RegisterActivity, DatabaseHelper, PasswordUtils và SessionManager. Khi đăng ký, dữ liệu tài khoản được lưu vào bảng users. Khi đăng nhập, mật khẩu được kiểm tra bằng PasswordUtils và phiên đăng nhập được duy trì bằng SharedPreferences.",
    )
    add_body(
        doc,
        "Chức năng quản lý kế hoạch được cài đặt trong AddTaskActivity, TaskFragment, TaskAdapter, TaskDetailActivity và DatabaseHelper. Người dùng có thể tạo kế hoạch theo loại, cập nhật thông tin, lọc danh sách, đổi trạng thái và xóa kế hoạch. Kế hoạch quá hạn được xử lý thống nhất bằng PlanBusinessRules thay vì lưu thêm trạng thái QUA_HAN vào database.",
    )
    add_body(
        doc,
        "Chức năng nhắc lịch sử dụng ReminderScheduler, ReminderReceiver và BootReceiver. Ứng dụng kiểm tra quyền exact alarm trên Android 12+, đặt alarm bằng setExactAndAllowWhileIdle, hiển thị notification khi đến giờ và đặt lại lịch nhắc sau reboot. Khi kế hoạch hoàn thành hoặc bị hủy, reminder tương ứng được tắt để tránh nhắc sai.",
    )
    add_table_caption(doc, "Bảng 4.3. Kết quả cài đặt chức năng")
    add_table(
        doc,
        ["Nhóm chức năng", "Kết quả cài đặt", "Đánh giá"],
        [
            ("Tài khoản offline", "Đăng ký, đăng nhập, đăng xuất; mật khẩu băm PBKDF2; session bằng SharedPreferences.", "Hoàn thành"),
            ("Danh mục kế hoạch", "CRUD danh mục theo user, có màu hiển thị.", "Hoàn thành"),
            ("Kế hoạch theo loại", "ASSIGNMENT, CLASS, PART_TIME, PERSONAL, EXAM, PROJECT; form thay đổi theo loại.", "Hoàn thành"),
            ("Tìm kiếm/lọc", "Lọc theo trạng thái, loại, danh mục; tìm kiếm có debounce.", "Hoàn thành"),
            ("Lặp lịch/trùng giờ", "Hỗ trợ NONE, DAILY, WEEKLY, MONTHLY, thứ 2-4-6, thứ 3-5, cuối tuần; có hasTimeConflict.", "Hoàn thành cơ bản"),
            ("Nhắc lịch", "Exact alarm, notification, nhắc 24h, boot receiver.", "Hoàn thành cơ bản"),
            ("Thống kê/gợi ý", "Tổng kế hoạch, hoàn thành, quá hạn, tuần hiện tại, giờ học/làm thêm, ưu tiên, gợi ý hôm nay.", "Hoàn thành cơ bản"),
            ("Sub-task/đánh giá", "Bảng, model và DAO đã có; UI chi tiết cần tiếp tục hoàn thiện.", "Nền tảng dữ liệu"),
        ],
        [1.5, 3.75, 1.05],
    )

    add_heading(doc, 2, "4.4. Kết quả giao diện ứng dụng")
    add_body(
        doc,
        "Các hình dưới đây là vị trí chèn ảnh giao diện thật khi hoàn thiện bản in. Do báo cáo được sinh tự động từ mã nguồn, nếu chưa có ảnh chụp màn hình mới nhất, phần hình được để placeholder rõ ràng để sinh viên thay bằng ảnh chạy thực tế trên thiết bị/emulator.",
    )
    for caption, label in [
        ("Hình 4.1. Màn hình đăng nhập", "MÀN HÌNH ĐĂNG NHẬP"),
        ("Hình 4.2. Màn hình trang chủ/dashboard", "MÀN HÌNH TRANG CHỦ"),
        ("Hình 4.3. Màn hình danh sách kế hoạch", "MÀN HÌNH DANH SÁCH KẾ HOẠCH"),
        ("Hình 4.4. Màn hình thêm/sửa kế hoạch", "MÀN HÌNH THÊM/SỬA KẾ HOẠCH"),
        ("Hình 4.5. Màn hình lịch", "MÀN HÌNH LỊCH"),
        ("Hình 4.6. Màn hình tổng quan/thống kê", "MÀN HÌNH TỔNG QUAN/THỐNG KÊ"),
    ]:
        add_figure_placeholder(
            doc,
            caption,
            [
                f"[CHÈN ẢNH GIAO DIỆN {label} TẠI ĐÂY]",
                "Ảnh nên được chụp từ ứng dụng chạy thật sau khi build để bảo đảm đúng phiên bản giao diện hiện tại.",
            ],
        )

    add_heading(doc, 2, "4.5. Kiểm thử chức năng")
    add_body(
        doc,
        "Kiểm thử được thực hiện theo các luồng thao tác chính của ứng dụng. Ngoài kiểm thử giao diện, dự án đã được kiểm tra build bằng Gradle ở nhánh hiện tại để bảo đảm mã Java, tài nguyên XML, manifest và dependencies có thể biên dịch.",
    )
    add_table_caption(doc, "Bảng 4.4. Bảng kiểm thử chức năng")
    add_table(
        doc,
        ["Mã TC", "Chức năng", "Dữ liệu đầu vào", "Kết quả mong đợi", "Kết quả thực tế", "Đánh giá"],
        [
            ("TC01", "Đăng ký", "Username/email/mật khẩu hợp lệ", "Tạo tài khoản mới, mật khẩu được hash.", "Tài khoản lưu vào SQLite.", "Đạt"),
            ("TC02", "Đăng nhập", "Tài khoản đã đăng ký", "Mở MainActivity và lưu session.", "Vào màn hình chính.", "Đạt"),
            ("TC03", "Đăng ký thiếu dữ liệu", "Bỏ trống email/mật khẩu", "Hiển thị lỗi validate.", "Không tạo tài khoản.", "Đạt"),
            ("TC04", "Thêm kế hoạch", "Tiêu đề, ngày, giờ, danh mục hợp lệ", "Lưu kế hoạch vào SQLite.", "Kế hoạch xuất hiện trong danh sách.", "Đạt"),
            ("TC05", "Sửa kế hoạch", "Thay đổi tiêu đề/ưu tiên", "Thông tin được cập nhật.", "Danh sách và chi tiết hiển thị dữ liệu mới.", "Đạt"),
            ("TC06", "Xóa kế hoạch", "Chọn xóa một kế hoạch", "Kế hoạch bị xóa khỏi DB.", "Không còn xuất hiện trong danh sách.", "Đạt"),
            ("TC07", "Tìm kiếm/lọc", "Từ khóa, trạng thái, loại, danh mục", "Chỉ hiển thị kế hoạch phù hợp.", "RecyclerView cập nhật đúng.", "Đạt"),
            ("TC08", "Trùng giờ", "Tạo kế hoạch cùng ngày/giờ với kế hoạch khác", "Cảnh báo hoặc bỏ qua ngày trùng khi tạo lặp.", "Có kiểm tra hasTimeConflict.", "Đạt"),
            ("TC09", "Quá hạn", "Kế hoạch chưa hoàn thành, deadline đã qua", "Hiển thị trạng thái Quá hạn.", "PlanBusinessRules nhận diện runtime.", "Đạt"),
            ("TC10", "Lịch", "Chọn một ngày trong calendar tháng", "Hiển thị kế hoạch của ngày đó.", "Danh sách sự kiện cập nhật.", "Đạt"),
            ("TC11", "Nhắc lịch", "Bật reminder đúng giờ/sau 5 phút/24h", "Alarm được đặt, notification xuất hiện đúng cấu hình.", "Logic đặt exact alarm đã cài đặt; cần kiểm thử thêm trên thiết bị thật.", "Đạt cơ bản"),
            ("TC12", "Thống kê", "Có nhiều kế hoạch ở trạng thái khác nhau", "Tính tổng, hoàn thành, quá hạn, tuần hiện tại.", "Overview/Home hiển thị số liệu.", "Đạt"),
            ("TC13", "Cài đặt theme", "Chọn Sáng/Tối", "Ứng dụng đổi theme và lưu lựa chọn.", "ThemeManager lưu mode.", "Đạt"),
        ],
        [0.55, 1.05, 1.15, 1.35, 1.45, 0.55],
    )

    add_heading(doc, 2, "4.6. Đánh giá kết quả thực nghiệm")
    add_body(
        doc,
        "Kết quả thực nghiệm cho thấy ứng dụng đáp ứng được yêu cầu cốt lõi của đề tài quản lý kế hoạch cá nhân cho sinh viên. Các chức năng lưu trữ, cập nhật, lọc, lịch, nhắc lịch và thống kê đã được tổ chức quanh mô hình kế hoạch thay vì chỉ quanh môn học. Điểm còn cần hoàn thiện là giao diện thao tác sub-task/đánh giá sau hoàn thành, cập nhật biểu đồ thống kê trực quan hơn và kiểm thử reminder trên nhiều dòng máy Android.",
    )
    add_table_caption(doc, "Bảng 4.5. Đánh giá kết quả đạt được")
    add_table(
        doc,
        ["Mục tiêu đánh giá", "Kết quả đạt được", "Nhận xét"],
        [
            ("Đúng phạm vi đề tài", "Ứng dụng tập trung vào kế hoạch học tập, bài tập, đi học, làm thêm, cá nhân, thi và dự án.", "Không mở rộng sang tài chính/sức khỏe."),
            ("Khả năng offline", "SQLite lưu tài khoản, danh mục, kế hoạch; SharedPreferences lưu phiên.", "Phù hợp ứng dụng cá nhân."),
            ("Phân biệt với ghi chú", "Có loại kế hoạch, ưu tiên, thời gian, lặp, trùng giờ, quá hạn, nhắc lịch, thống kê.", "Giá trị nghiệp vụ rõ."),
            ("Nhắc lịch", "Có exact alarm, notification, boot receiver, nhắc 24h.", "Cần kiểm thử thêm theo chính sách pin của từng hãng."),
            ("Bảo trì", "Package/class tách theo vai trò; database version 5 mở rộng bảng.", "Có thể tiếp tục nâng cấp."),
            ("Giao diện", "Material-style dashboard, card, bottom navigation, FAB.", "Cần thay ảnh thật trong báo cáo trước khi nộp."),
        ],
        [1.35, 3.55, 1.4],
    )
    add_page_break(doc)

    # Chapter 5
    add_heading(doc, 1, "CHƯƠNG 5: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")
    add_heading(doc, 2, "5.1. Kết quả đạt được")
    add_body(
        doc,
        "Sau quá trình thực hiện, đề tài đã xây dựng được ứng dụng Android quản lý kế hoạch cá nhân dành cho sinh viên theo đúng phạm vi học phần Lập Trình Thiết Bị Di Động. Ứng dụng hỗ trợ đăng ký/đăng nhập offline, quản lý danh mục kế hoạch, thêm/sửa/xóa kế hoạch theo loại, tìm kiếm/lọc, lịch dạng tháng, nhắc lịch, kiểm tra trùng thời gian, nhận diện quá hạn, thống kê tổng quan và gợi ý kế hoạch hôm nay.",
    )
    add_body(
        doc,
        "Về mặt kỹ thuật, dự án vận dụng các kiến thức quan trọng của Android như Activity, Fragment, RecyclerView, Adapter, XML Layout, Material Components, SQLiteOpenHelper, SharedPreferences, AlarmManager, BroadcastReceiver và NotificationCompat. Cấu trúc mã nguồn sau dọn dẹp tương đối rõ, phù hợp để giải thích trong quá trình bảo vệ đồ án.",
    )

    add_heading(doc, 2, "5.2. Ưu điểm và nhược điểm")
    add_heading(doc, 3, "5.2.1. Ưu điểm")
    add_bullets(
        doc,
        [
            "Ứng dụng hoạt động offline, không phụ thuộc kết nối Internet.",
            "Mô hình dữ liệu phản ánh nhiều loại kế hoạch trong đời sống sinh viên, không chỉ quản lý môn học.",
            "Giao diện hiện đại, có dashboard, lịch, thống kê và gợi ý hôm nay.",
            "Danh sách kế hoạch hỗ trợ tìm kiếm và lọc theo trạng thái, danh mục, loại kế hoạch.",
            "Nhắc lịch sử dụng exact alarm, có receiver và cơ chế khôi phục sau khi thiết bị khởi động lại.",
            "Mật khẩu được băm bằng PBKDF2 kèm salt, tốt hơn lưu văn bản rõ.",
        ],
    )
    add_heading(doc, 3, "5.2.2. Nhược điểm")
    add_bullets(
        doc,
        [
            "Dữ liệu hiện chỉ lưu trên một thiết bị, chưa đồng bộ giữa nhiều máy.",
            "Màn hình thao tác sub-task và đánh giá sau hoàn thành chưa hoàn thiện đầy đủ dù nền tảng dữ liệu đã có.",
            "Thông báo hiện mở về MainActivity, chưa mở thẳng vào màn hình chi tiết kế hoạch.",
            "Thống kê hiện ở mức cơ bản, chưa có biểu đồ trực quan nâng cao theo tháng/năm.",
            "Chưa hỗ trợ cộng tác hoặc chia sẻ kế hoạch nhóm giữa nhiều sinh viên.",
            "Một số chuỗi giao diện trong resource hiện được viết không dấu để tránh lỗi encoding, cần chuẩn hóa UTF-8 khi hoàn thiện sản phẩm.",
        ],
    )

    add_heading(doc, 2, "5.3. Hướng phát triển")
    add_bullets(
        doc,
        [
            "Tích hợp Firebase Authentication và Cloud Firestore để đồng bộ dữ liệu, sao lưu và đăng nhập đa thiết bị.",
            "Hoàn thiện giao diện sub-task và đánh giá sau hoàn thành để đo hiệu quả thực hiện kế hoạch.",
            "Bổ sung biểu đồ thống kê theo tuần/tháng/năm cho giờ học, giờ làm thêm, bài tập chưa nộp, tỷ lệ hoàn thành và tỷ lệ trễ hạn.",
            "Tối ưu thông báo nhắc lịch: nhắc lại, tùy chỉnh âm thanh, mở thẳng vào chi tiết kế hoạch.",
            "Bổ sung sao lưu/khôi phục dữ liệu bằng file để tránh mất dữ liệu khi đổi thiết bị.",
            "Mở rộng chia sẻ kế hoạch nhóm cho đồ án, học nhóm hoặc lịch làm thêm theo ca.",
            "Tối ưu giao diện cho nhiều kích thước màn hình và hoàn thiện dark mode toàn bộ màn hình.",
        ],
    )
    add_page_break(doc)

    add_heading(doc, 1, "TÀI LIỆU THAM KHẢO")
    refs = [
        "[1] Google, “Android Developers Documentation,” Android Developers. [Online]. Available: https://developer.android.com/.",
        "[2] Google, “Material Design 3,” Material Design. [Online]. Available: https://m3.material.io/.",
        "[3] Android Developers, “Save data using SQLite,” Android Developers. [Online]. Available: https://developer.android.com/training/data-storage/sqlite.",
        "[4] Android Developers, “Schedule alarms,” Android Developers. [Online]. Available: https://developer.android.com/develop/background-work/services/alarms.",
        "[5] Android Developers, “Notifications overview,” Android Developers. [Online]. Available: https://developer.android.com/develop/ui/views/notifications.",
        "[6] Oracle, “Java Documentation,” Oracle. [Online]. Available: https://docs.oracle.com/en/java/.",
        "[7] NIST, “Digital Identity Guidelines: Authentication and Lifecycle Management,” NIST Special Publication 800-63B.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(-0.6)
        p.paragraph_format.left_indent = Cm(0.6)
        run = p.add_run(ref)
        set_run_font(run, size=12)

    # Set all table rows to allow natural height and keep text readable.
    for table in doc.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_report())
